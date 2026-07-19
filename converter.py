from __future__ import annotations

import gzip
import html
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz
import imageio_ffmpeg
from docx import Document
from PIL import Image, ImageSequence, UnidentifiedImageError


class ConversionError(Exception):
    """A safe, user-facing conversion error."""


FORMAT_INFO = {
    "pdf": {"label": "PDF", "category": "documents"},
    "docx": {"label": "Word", "category": "documents"},
    "doc": {"label": "Word (legacy)", "category": "documents"},
    "txt": {"label": "Text", "category": "documents"},
    "md": {"label": "Markdown", "category": "documents"},
    "rtf": {"label": "Rich Text", "category": "documents"},
    "html": {"label": "HTML", "category": "documents"},
    "odt": {"label": "OpenDocument", "category": "documents"},
    "ppt": {"label": "PowerPoint (legacy)", "category": "documents"},
    "pptx": {"label": "PowerPoint", "category": "documents"},
    "xls": {"label": "Excel (legacy)", "category": "documents"},
    "xlsx": {"label": "Excel", "category": "documents"},
    "png": {"label": "PNG", "category": "images"},
    "jpg": {"label": "JPG", "category": "images"},
    "webp": {"label": "WebP", "category": "images"},
    "bmp": {"label": "BMP", "category": "images"},
    "tiff": {"label": "TIFF", "category": "images"},
    "ico": {"label": "ICO", "category": "images"},
    "gif": {"label": "GIF", "category": "animation"},
    "tgs": {"label": "TGS", "category": "animation"},
    "json": {"label": "Lottie JSON", "category": "animation"},
    "mp4": {"label": "MP4", "category": "video"},
    "webm": {"label": "WebM", "category": "video"},
    "mov": {"label": "MOV", "category": "video"},
    "mkv": {"label": "MKV", "category": "video"},
    "avi": {"label": "AVI", "category": "video"},
    "m4v": {"label": "M4V", "category": "video"},
    "mpeg": {"label": "MPEG", "category": "video"},
    "3gp": {"label": "3GP", "category": "video"},
    "mp3": {"label": "MP3", "category": "audio"},
    "wav": {"label": "WAV", "category": "audio"},
    "ogg": {"label": "OGG", "category": "audio"},
    "flac": {"label": "FLAC", "category": "audio"},
    "m4a": {"label": "M4A", "category": "audio"},
    "aac": {"label": "AAC", "category": "audio"},
    "wma": {"label": "WMA", "category": "audio"},
    "opus": {"label": "Opus", "category": "audio"},
}

ALIASES = {"jpeg": "jpg", "jpe": "jpg", "tif": "tiff", "htm": "html", "mpg": "mpeg"}

IMAGE_INPUTS = {"png", "jpg", "webp", "bmp", "tiff", "ico"}
IMAGE_OUTPUTS = ["png", "jpg", "webp", "gif", "bmp", "tiff", "ico", "pdf"]
VIDEO_INPUTS = {"mp4", "webm", "mov", "mkv", "avi", "m4v", "mpeg", "3gp"}
VIDEO_OUTPUTS = [
    "mp4", "webm", "mov", "mkv", "avi", "m4v", "mpeg", "3gp",
    "gif", "png", "jpg", "webp",
    "mp3", "wav", "ogg", "flac", "m4a", "aac", "opus", "wma",
]
AUDIO_INPUTS = {"mp3", "wav", "ogg", "flac", "m4a", "aac", "wma", "opus"}
AUDIO_OUTPUTS = ["mp3", "wav", "ogg", "flac", "m4a", "aac", "opus", "wma"]
OFFICE_TEXT_INPUTS = {"doc", "odt", "rtf"}
OFFICE_SHEET_INPUTS = {"xls", "xlsx"}
OFFICE_SLIDE_INPUTS = {"ppt", "pptx"}


def _clean_basename(filename: str) -> str:
    name = filename.replace("\\", "/").rsplit("/", 1)[-1]
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip(" .")
    return name or "untitled"


def output_filename(original_name: str, target: str, used_names: set[str]) -> str:
    clean_name = _clean_basename(original_name)
    stem = Path(clean_name).stem.strip(" .") or "untitled"
    candidate = f"{stem}.{target}"
    counter = 2
    while candidate.casefold() in used_names:
        candidate = f"{stem} ({counter}).{target}"
        counter += 1
    used_names.add(candidate.casefold())
    return candidate


class ConversionManager:
    def __init__(self) -> None:
        self.ffmpeg = self._find_ffmpeg()
        self.libreoffice = self._find_libreoffice()
        self.conversions = self._build_conversions()

    @staticmethod
    def _find_ffmpeg() -> str | None:
        system_ffmpeg = shutil.which("ffmpeg")
        if system_ffmpeg:
            return system_ffmpeg
        try:
            bundled = imageio_ffmpeg.get_ffmpeg_exe()
            return bundled if Path(bundled).exists() else None
        except Exception:
            return None

    @staticmethod
    def _find_libreoffice() -> str | None:
        for executable in ("soffice", "libreoffice"):
            found = shutil.which(executable)
            if found:
                return found
        if os.name == "nt":
            for candidate in (
                Path(os.environ.get("ProgramFiles", "")) / "LibreOffice/program/soffice.exe",
                Path(os.environ.get("ProgramFiles(x86)", "")) / "LibreOffice/program/soffice.exe",
            ):
                if candidate.exists():
                    return str(candidate)
        return None

    def _build_conversions(self) -> dict[str, list[str]]:
        conversions: dict[str, list[str]] = {
            "pdf": ["pdf", "docx", "txt"],
            "docx": ["pdf", "txt", "html"],
            "txt": ["pdf", "docx", "txt"],
            "md": ["pdf", "docx", "txt"],
            "html": ["pdf", "docx", "txt"],
            "rtf": ["pdf", "docx", "txt"],
            "gif": IMAGE_OUTPUTS.copy(),
            "tgs": ["json"],
            "json": ["tgs"],
        }
        for source in IMAGE_INPUTS:
            conversions[source] = IMAGE_OUTPUTS.copy()
        if self.ffmpeg:
            for source in IMAGE_INPUTS:
                conversions[source] += ["mp4", "webm"]
            for source in VIDEO_INPUTS:
                conversions[source] = VIDEO_OUTPUTS.copy()
            for source in AUDIO_INPUTS:
                conversions[source] = AUDIO_OUTPUTS.copy()
            conversions["gif"] += ["mp4", "webm"]
        if self.libreoffice:
            for source in OFFICE_TEXT_INPUTS:
                conversions[source] = ["pdf", "docx", "txt"]
            for source in OFFICE_SHEET_INPUTS | OFFICE_SLIDE_INPUTS:
                conversions[source] = ["pdf"]
        return conversions

    def capabilities(self) -> dict:
        categories = [
            {"id": "documents", "label": "Documents", "description": "PDF, Word & text"},
            {"id": "images", "label": "Images", "description": "PNG, JPG & WebP"},
            {"id": "audio", "label": "Audio", "description": "MP3, WAV & more"},
            {"id": "video", "label": "Video", "description": "MP4, WebM & more"},
            {"id": "animation", "label": "Animation", "description": "GIF, TGS & Lottie"},
        ]
        formats = {
            key: {**FORMAT_INFO[key], "targets": targets}
            for key, targets in self.conversions.items()
        }
        return {
            "categories": categories,
            "formats": formats,
            "accept": [f".{extension}" for extension in sorted(formats)],
            "engines": {
                "ffmpeg": bool(self.ffmpeg),
                "libreoffice": bool(self.libreoffice),
                "built_in_documents": True,
            },
            "limits": {"max_files": 50, "max_batch_bytes": 2 * 1024 * 1024 * 1024},
        }

    def detect_format(self, filename: str) -> str:
        suffix = Path(_clean_basename(filename)).suffix.lower().lstrip(".")
        suffix = ALIASES.get(suffix, suffix)
        if not suffix or suffix not in self.conversions:
            readable = suffix.upper() if suffix else "unknown"
            raise ConversionError(f"{readable} files are not supported by this local setup.")
        return suffix

    def validate_conversion(self, source: str, target: str) -> None:
        if source not in self.conversions or target not in self.conversions[source]:
            raise ConversionError(f"{source.upper()} cannot be converted to {target.upper()}.")

    def convert(self, source: Path, source_format: str, output: Path, target: str) -> str | None:
        self.validate_conversion(source_format, target)
        try:
            if source_format == "pdf":
                return self._convert_pdf(source, output, target)
            if source_format == "docx":
                return self._convert_docx(source, output, target)
            if source_format in {"txt", "md", "html", "rtf"}:
                return self._convert_text_document(source, source_format, output, target)
            if source_format in IMAGE_INPUTS or source_format == "gif":
                if target in {"mp4", "webm"}:
                    animated = self._is_animated_image(source)
                    if animated:
                        self._convert_media(source, output, target, mute=True)
                    else:
                        self._convert_still_image_to_media(source, output, target)
                    return None if animated else "Static images become a three-second video."
                return self._convert_image(source, output, target)
            if source_format in VIDEO_INPUTS:
                self._convert_media(source, output, target)
                return "The first video frame was exported." if target in {"png", "jpg", "webp"} else None
            if source_format in AUDIO_INPUTS:
                self._convert_media(source, output, target)
                return None
            if source_format == "tgs":
                self._tgs_to_json(source, output)
                return "TGS is exported as editable Lottie JSON."
            if source_format == "json":
                self._json_to_tgs(source, output)
                return None
            if source_format in OFFICE_TEXT_INPUTS | OFFICE_SHEET_INPUTS | OFFICE_SLIDE_INPUTS:
                return self._convert_with_libreoffice(source, output, target)
        except ConversionError:
            raise
        except (OSError, ValueError, UnidentifiedImageError, fitz.FileDataError) as exc:
            raise ConversionError("The file is damaged or its contents do not match its extension.") from exc
        raise ConversionError("No converter is available for this format pair.")

    @staticmethod
    def _is_animated_image(source: Path) -> bool:
        with Image.open(source) as image:
            return getattr(image, "n_frames", 1) > 1

    def _convert_still_image_to_media(self, source: Path, output: Path, target: str) -> None:
        # Normalizing through PNG makes FFmpeg looping consistent for formats
        # such as ICO whose native demuxer does not accept the image2 loop option.
        with tempfile.TemporaryDirectory(prefix="localconvert-frame-") as directory:
            normalized = Path(directory) / "frame.png"
            with Image.open(source) as image:
                rgba = image.convert("RGBA")
                background = Image.new("RGB", rgba.size, "white")
                background.paste(rgba, mask=rgba.getchannel("A"))
                background.save(normalized, "PNG")
            self._convert_media(normalized, output, target, mute=True, still_image=True)

    @staticmethod
    def _convert_image(source: Path, output: Path, target: str) -> str | None:
        if target == "pdf":
            with Image.open(source) as image:
                frame = image.convert("RGB")
                frame.save(output, "PDF", resolution=150)
            return "Animated images use their first frame when exported to PDF."

        with Image.open(source) as image:
            frames = [frame.copy() for frame in ImageSequence.Iterator(image)]
            durations = [frame.info.get("duration", image.info.get("duration", 100)) for frame in frames]
            loop = image.info.get("loop", 0)
            animated = len(frames) > 1

            if target in {"gif", "webp"} and animated:
                converted = []
                for frame in frames:
                    converted.append(frame.convert("RGBA" if target == "webp" else "P"))
                save_args = {
                    "save_all": True,
                    "append_images": converted[1:],
                    "duration": durations,
                    "loop": loop,
                }
                if target == "webp":
                    save_args.update({"quality": 90, "method": 6})
                else:
                    save_args.update({"optimize": True, "disposal": 2})
                converted[0].save(output, target.upper(), **save_args)
                return None

            frame = frames[0]
            if target == "jpg":
                if frame.mode in {"RGBA", "LA", "P"}:
                    rgba = frame.convert("RGBA")
                    background = Image.new("RGB", rgba.size, "white")
                    background.paste(rgba, mask=rgba.getchannel("A"))
                    frame = background
                else:
                    frame = frame.convert("RGB")
                frame.save(output, "JPEG", quality=92, optimize=True)
            elif target == "png":
                frame.save(output, "PNG", optimize=True)
            elif target == "webp":
                frame.save(output, "WEBP", quality=90, method=6)
            elif target == "bmp":
                frame.convert("RGB").save(output, "BMP")
            elif target == "tiff":
                frame.save(output, "TIFF", compression="tiff_deflate")
            elif target == "ico":
                icon_source = frame.convert("RGBA")
                icon_source.thumbnail((256, 256), Image.Resampling.LANCZOS)
                icon_canvas = Image.new("RGBA", (256, 256), (0, 0, 0, 0))
                offset = ((256 - icon_source.width) // 2, (256 - icon_source.height) // 2)
                icon_canvas.paste(icon_source, offset, icon_source)
                icon_canvas.save(output, "ICO", sizes=[(16, 16), (32, 32), (48, 48), (64, 64), (128, 128), (256, 256)])
            elif target == "gif":
                frame.convert("P").save(output, "GIF", optimize=True)
            else:
                raise ConversionError(f"Image conversion to {target.upper()} is unavailable.")
        return "Animated images use their first frame for this output format." if animated else None

    @staticmethod
    def _convert_pdf(source: Path, output: Path, target: str) -> str | None:
        with fitz.open(source) as pdf:
            if pdf.needs_pass:
                raise ConversionError("Password-protected PDFs must be unlocked before conversion.")
            if target == "pdf":
                pdf.save(output, garbage=4, deflate=True, clean=True)
                return "The PDF was cleaned and optimized locally."
            if target == "txt":
                text = "\n\n".join(page.get_text("text").rstrip() for page in pdf)
                output.write_text(text, encoding="utf-8")
                return "Scanned pages need OCR before their text can be extracted."
            if target == "docx":
                document = Document()
                for page_number, page in enumerate(pdf):
                    blocks = page.get_text("blocks", sort=True)
                    for block in blocks:
                        block_text = block[4].strip()
                        if block_text:
                            document.add_paragraph(block_text)
                    if page_number < len(pdf) - 1:
                        document.add_page_break()
                document.save(output)
                return "Text is preserved; complex PDF layout, forms, and columns may shift in Word."
        raise ConversionError(f"PDF conversion to {target.upper()} is unavailable.")

    @staticmethod
    def _docx_html(source: Path) -> str:
        document = Document(source)
        chunks: list[str] = []
        for paragraph in document.paragraphs:
            text = html.escape(paragraph.text)
            style = (paragraph.style.name if paragraph.style else "").lower()
            if style.startswith("heading"):
                match = re.search(r"(\d+)", style)
                level = min(int(match.group(1)), 3) if match else 2
                chunks.append(f"<h{level}>{text}</h{level}>")
            elif text:
                chunks.append(f"<p>{text}</p>")
            else:
                chunks.append("<p><br></p>")
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = "".join(f"<td>{html.escape(cell.text)}</td>" for cell in row.cells)
                rows.append(f"<tr>{cells}</tr>")
            chunks.append(f"<table>{''.join(rows)}</table>")
        return "\n".join(chunks)

    def _convert_docx(self, source: Path, output: Path, target: str) -> str | None:
        body = self._docx_html(source)
        if target == "html":
            output.write_text(
                "<!doctype html><html><head><meta charset=\"utf-8\"><title>Converted document</title></head>"
                f"<body>{body}</body></html>",
                encoding="utf-8",
            )
            return "Text, headings, and tables are preserved; advanced Word layout may shift."
        if target == "txt":
            document = Document(source)
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            output.write_text("\n".join(parts), encoding="utf-8")
            return None
        if target == "pdf":
            self._write_pdf_from_html(body, output)
            return "Text, headings, and tables are preserved; advanced Word layout may shift in PDF."
        raise ConversionError(f"Word conversion to {target.upper()} is unavailable.")

    def _convert_text_document(self, source: Path, source_format: str, output: Path, target: str) -> str | None:
        raw = self._read_text(source)
        if source_format == "html":
            body = self._sanitize_html(raw)
            plain = self._html_to_text(body)
        elif source_format == "rtf":
            plain = self._rtf_to_text(raw)
            body = "".join(f"<p>{html.escape(line)}</p>" for line in plain.splitlines())
        elif source_format == "md":
            plain = raw
            body = self._markdown_to_html(raw)
        else:
            plain = raw
            body = "".join(f"<p>{html.escape(line) or '<br>'}</p>" for line in raw.splitlines())

        if target == "txt":
            output.write_text(plain, encoding="utf-8")
        elif target == "docx":
            document = Document()
            for line in plain.splitlines():
                document.add_paragraph(line)
            document.save(output)
        elif target == "pdf":
            self._write_pdf_from_html(body, output)
        else:
            raise ConversionError(f"Document conversion to {target.upper()} is unavailable.")
        return None

    @staticmethod
    def _read_text(source: Path) -> str:
        data = source.read_bytes()
        for encoding in ("utf-8-sig", "utf-16", "cp1252"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ConversionError("The document text encoding could not be read.")

    @staticmethod
    def _sanitize_html(raw: str) -> str:
        cleaned = re.sub(r"(?is)<(script|style|iframe|object).*?>.*?</\1>", "", raw)
        cleaned = re.sub(r"(?is)\son\w+\s*=\s*(['\"]).*?\1", "", cleaned)
        return cleaned

    @staticmethod
    def _html_to_text(raw: str) -> str:
        spaced = re.sub(r"(?i)</?(p|div|h[1-6]|li|tr|br)\b[^>]*>", "\n", raw)
        return html.unescape(re.sub(r"<[^>]+>", "", spaced)).strip()

    @staticmethod
    def _rtf_to_text(raw: str) -> str:
        text = re.sub(r"\\par[d]?\b", "\n", raw)
        text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
        text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
        return text.replace("{", "").replace("}", "").strip()

    @staticmethod
    def _markdown_to_html(raw: str) -> str:
        chunks = []
        for line in raw.splitlines():
            escaped = html.escape(line)
            heading = re.match(r"^(#{1,3})\s+(.*)$", escaped)
            if heading:
                level = len(heading.group(1))
                chunks.append(f"<h{level}>{heading.group(2)}</h{level}>")
            elif escaped.startswith(("- ", "* ")):
                chunks.append(f"<p>• {escaped[2:]}</p>")
            else:
                chunks.append(f"<p>{escaped or '<br>'}</p>")
        return "".join(chunks)

    @staticmethod
    def _write_pdf_from_html(body: str, output: Path) -> None:
        css = """
            body { font-family: sans-serif; font-size: 10.5pt; line-height: 1.45; color: #17211b; }
            h1 { font-size: 22pt; margin: 0 0 12pt; } h2 { font-size: 17pt; } h3 { font-size: 13pt; }
            p { margin: 0 0 8pt; } table { border-collapse: collapse; width: 100%; margin: 8pt 0; }
            td, th { border: 0.6pt solid #aeb8b1; padding: 5pt; }
        """
        story = fitz.Story(html=f"<body>{body}</body>", user_css=css)
        writer = fitz.DocumentWriter(str(output))
        page_rect = fitz.paper_rect("a4")
        content_rect = page_rect + (54, 54, -54, -54)
        more = 1
        while more:
            device = writer.begin_page(page_rect)
            more, _filled = story.place(content_rect)
            story.draw(device)
            writer.end_page()
        writer.close()

    def _convert_media(self, source: Path, output: Path, target: str, mute: bool = False, still_image: bool = False) -> None:
        if not self.ffmpeg:
            raise ConversionError("FFmpeg is not available in this local setup.")
        command = [self.ffmpeg, "-hide_banner", "-loglevel", "error", "-y"]
        if still_image:
            command += ["-loop", "1", "-framerate", "30"]
        command += ["-i", str(source)]
        if target == "mp4":
            command += ["-vf", "scale=trunc(iw/2)*2:trunc(ih/2)*2", "-c:v", "libx264", "-preset", "medium", "-crf", "22", "-pix_fmt", "yuv420p"]
            command += ["-an"] if mute else ["-c:a", "aac", "-b:a", "192k"]
            command += ["-movflags", "+faststart"]
        elif target == "webm":
            command += ["-c:v", "libvpx-vp9", "-crf", "31", "-b:v", "0", "-pix_fmt", "yuv420p"]
            command += ["-an"] if mute else ["-c:a", "libopus", "-b:a", "128k"]
        elif target == "mov":
            command += ["-c:v", "libx264", "-crf", "22", "-pix_fmt", "yuv420p", "-c:a", "aac", "-b:a", "192k"]
        elif target == "mkv":
            command += ["-c:v", "libx264", "-crf", "22", "-c:a", "aac", "-b:a", "192k"]
        elif target == "avi":
            command += ["-c:v", "mpeg4", "-q:v", "4", "-c:a", "libmp3lame", "-b:a", "192k"]
        elif target == "m4v":
            command += ["-c:v", "mpeg4", "-q:v", "4", "-an", "-f", "m4v"]
        elif target == "mpeg":
            command += ["-c:v", "mpeg2video", "-q:v", "4", "-c:a", "mp2", "-b:a", "192k", "-f", "mpeg"]
        elif target == "3gp":
            command += ["-vf", "scale=trunc(iw/2)*2:trunc(ih/2)*2", "-c:v", "libx264", "-profile:v", "baseline", "-level", "3.0", "-pix_fmt", "yuv420p", "-c:a", "aac", "-ar", "44100", "-ac", "2", "-f", "3gp"]
        elif target == "gif":
            command += ["-vf", "fps=15,scale=960:-2:force_original_aspect_ratio=decrease:flags=lanczos", "-loop", "0"]
        elif target == "png":
            command += ["-frames:v", "1", "-c:v", "png"]
        elif target == "jpg":
            command += ["-frames:v", "1", "-q:v", "2"]
        elif target == "webp":
            command += ["-frames:v", "1", "-c:v", "libwebp", "-quality", "90"]
        elif target == "mp3":
            command += ["-vn", "-c:a", "libmp3lame", "-q:a", "2"]
        elif target == "wav":
            command += ["-vn", "-c:a", "pcm_s16le"]
        elif target == "ogg":
            command += ["-vn", "-c:a", "libvorbis", "-q:a", "5"]
        elif target == "flac":
            command += ["-vn", "-c:a", "flac"]
        elif target in {"m4a", "aac"}:
            command += ["-vn", "-c:a", "aac", "-b:a", "192k"]
        elif target == "opus":
            command += ["-vn", "-c:a", "libopus", "-b:a", "128k"]
        elif target == "wma":
            command += ["-vn", "-c:a", "wmav2", "-b:a", "192k"]
        else:
            raise ConversionError(f"Media conversion to {target.upper()} is unavailable.")
        if still_image:
            command += ["-t", "3"]
        command.append(str(output))
        creation_flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=60 * 60,
                creationflags=creation_flags,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise ConversionError("Conversion exceeded the one-hour safety limit.") from exc
        if completed.returncode != 0:
            detail = completed.stderr.strip().splitlines()
            message = detail[-1] if detail else "FFmpeg could not read this media file."
            raise ConversionError(f"Media conversion failed: {message[:240]}")

    @staticmethod
    def _tgs_to_json(source: Path, output: Path) -> None:
        try:
            with gzip.open(source, "rt", encoding="utf-8") as compressed:
                data = json.load(compressed)
        except (gzip.BadGzipFile, json.JSONDecodeError, UnicodeDecodeError) as exc:
            raise ConversionError("This is not a valid Telegram TGS animation.") from exc
        ConversionManager._validate_lottie(data)
        output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    @staticmethod
    def _json_to_tgs(source: Path, output: Path) -> None:
        try:
            data = json.loads(source.read_text(encoding="utf-8-sig"))
        except (json.JSONDecodeError, UnicodeDecodeError) as exc:
            raise ConversionError("This is not valid Lottie JSON.") from exc
        ConversionManager._validate_lottie(data)
        compact = json.dumps(data, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
        with output.open("wb") as file_handle:
            with gzip.GzipFile(filename="", mode="wb", fileobj=file_handle, mtime=0) as compressed:
                compressed.write(compact)

    @staticmethod
    def _validate_lottie(data: object) -> None:
        required = {"v", "fr", "ip", "op", "w", "h", "layers"}
        if not isinstance(data, dict) or not required.issubset(data):
            raise ConversionError("The JSON file is not a complete Lottie animation.")

    def _convert_with_libreoffice(self, source: Path, output: Path, target: str) -> str | None:
        if not self.libreoffice:
            raise ConversionError("LibreOffice is required for this Office format.")
        with tempfile.TemporaryDirectory(prefix="localconvert-office-") as output_dir:
            command = [
                self.libreoffice,
                "--headless",
                "--convert-to",
                target,
                "--outdir",
                output_dir,
                str(source),
            ]
            creation_flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
            completed = subprocess.run(command, capture_output=True, text=True, timeout=600, creationflags=creation_flags)
            candidates = list(Path(output_dir).glob(f"*.{target}"))
            if completed.returncode != 0 or not candidates:
                raise ConversionError("LibreOffice could not convert this document.")
            shutil.move(candidates[0], output)
        return "LibreOffice handled this legacy Office conversion locally."
